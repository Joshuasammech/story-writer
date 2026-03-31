"""
Story Writer Bot — Web Interface
Flask app with simple username/password login + SSE streaming.
"""

import io
import json
import os
import re
from functools import wraps

from dotenv import load_dotenv
load_dotenv()

import csv

import anthropic
import requests
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import (
    Flask, Response, redirect, request, send_file,
    session, stream_with_context, url_for,
)

# ── App setup ──────────────────────────────────────────────────────────────────

app = Flask(__name__, static_folder="static")
app.secret_key = os.environ.get("SECRET_KEY", os.urandom(32))

LOGIN_USERNAME = os.environ.get("LOGIN_USERNAME", "admin")
LOGIN_PASSWORD = os.environ.get("LOGIN_PASSWORD", "changeme")

_api_key = os.environ.get("ANTHROPIC_API_KEY", "")
if _api_key:
    print(f"[startup] ANTHROPIC_API_KEY loaded ({len(_api_key)} chars, starts with {_api_key[:12]}…)")
else:
    print("[startup] WARNING: ANTHROPIC_API_KEY is NOT set — story generation will fail")

# ── Story prompt ───────────────────────────────────────────────────────────────

SYSTEM_PROMPT = """\
You are a transformation case study writer. Write a Leadership Transformation Case Study from the report. \
Write with warmth and precision — like someone who deeply understands people. \
Every fact must come from the report. Never invent details.

Use EXACTLY this format:

---
# [Full Name]
### [One line: what this transformation fundamentally is]

**[Key stat]** · **[Second stat]** · **[Key outcome]**

---

## 01 | Participant Profile
[2 sentences: who they are, their role or context, and why this journey matters.]

---

## 02 | Starting Conditions
[2 sentences on where they began — the challenges, gaps, or circumstances they faced before the transformation.]

---

## 03 | Impact & Breakthroughs
[3 sentences on the specific results, shifts, and breakthroughs achieved. Use concrete outcomes from the report.]

> "[One sentence capturing the defining moment or turning point.]"

---

## 04 | The Contrast
[2 sentences: compared to a conventional path or typical timeline, how was this different? How much faster, deeper, or more significant?]

---

## 05 | Next Steps
[2 sentences on where they are headed — what they are building, pursuing, or stepping into next.]

---

**CORE INSIGHT**
[2 sentences. The defining statement — what this transformation reveals about what is possible.]

---

Rules: no extra sections, no hollow phrases ("hard work", "passion", "incredible"), third person, 250–320 words total.\
"""

# ── Auth helpers ───────────────────────────────────────────────────────────────

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get("logged_in"):
            return redirect(url_for("login_page"))
        return f(*args, **kwargs)
    return decorated



# ── Helpers ────────────────────────────────────────────────────────────────────

def detect_url_type(url: str) -> str:
    """Return 'sheet' or 'doc' based on URL pattern."""
    if "spreadsheets" in url:
        return "sheet"
    return "doc"


def extract_doc_id(url_or_id: str) -> str:
    match = re.search(r"/d/([a-zA-Z0-9_-]+)", url_or_id)
    return match.group(1) if match else url_or_id.strip()


def _safe_get(url: str) -> requests.Response:
    """GET with retry on connection/timeout errors."""
    for attempt in range(2):
        try:
            return requests.get(url, timeout=25)
        except requests.exceptions.ConnectionError:
            if attempt == 1:
                raise RuntimeError(
                    "Could not reach Google. Check your internet connection or try again."
                )
        except requests.exceptions.Timeout:
            if attempt == 1:
                raise RuntimeError("Google request timed out. Try again.")


def fetch_google_sheet(url_or_id: str) -> tuple[str, str]:
    """Fetch a Google Sheet as CSV and return (title, formatted_text)."""
    sheet_id = extract_doc_id(url_or_id)
    # Preserve gid (tab) if present in the URL
    gid_match = re.search(r"[#&?]gid=(\d+)", url_or_id)
    gid_param = f"&gid={gid_match.group(1)}" if gid_match else ""
    export_url = (
        f"https://docs.google.com/spreadsheets/d/{sheet_id}"
        f"/export?format=csv{gid_param}"
    )
    resp = _safe_get(export_url)
    if resp.status_code == 403:
        raise RuntimeError(
            'Sheet access denied. Share it as "Anyone with the link can view".'
        )
    if resp.status_code == 404:
        raise RuntimeError("Sheet not found. Check the URL is correct.")
    if not resp.ok:
        raise RuntimeError(f"Failed to fetch sheet (HTTP {resp.status_code}).")

    # Convert CSV → readable text table
    lines = []
    reader = csv.reader(resp.text.splitlines())
    rows = list(reader)
    if not rows:
        raise RuntimeError("The sheet appears to be empty.")
    title = rows[0][0] if rows[0] else "Google Sheet"
    for row in rows:
        non_empty = [c.strip() for c in row if c.strip()]
        if non_empty:
            lines.append(" | ".join(non_empty))
    return title, "\n".join(lines)


def fetch_google_doc(url_or_id: str) -> tuple[str, str]:
    doc_id = extract_doc_id(url_or_id)
    export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=txt"
    resp = _safe_get(export_url)
    if resp.status_code == 403:
        raise RuntimeError(
            'Access denied. Share the doc as "Anyone with the link can view".'
        )
    if resp.status_code == 404:
        raise RuntimeError("Document not found. Check that the URL is correct.")
    if not resp.ok:
        raise RuntimeError(f"Failed to fetch document (HTTP {resp.status_code}).")
    text = resp.text.strip()
    lines = text.splitlines()
    title = next((l.strip() for l in lines if l.strip()), "Untitled Document")
    return title, text


def sse(event: str, data: str) -> str:
    payload = json.dumps({"data": data})
    return f"event: {event}\ndata: {payload}\n\n"


# ── Auth routes ────────────────────────────────────────────────────────────────

@app.route("/login", methods=["GET", "POST"])
def login_page():
    error = None
    if request.method == "POST":
        username = request.form.get("username", "")
        password = request.form.get("password", "")
        if username == LOGIN_USERNAME and password == LOGIN_PASSWORD:
            session["logged_in"] = True
            return redirect(url_for("index"))
        error = "Invalid username or password."
    with open(os.path.join(app.static_folder, "login.html"), encoding="utf-8") as f:
        html = f.read()
    if error:
        html = html.replace("<!--ERROR-->", f'<p class="error">{error}</p>')
    return html


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login_page"))


# ── App routes ─────────────────────────────────────────────────────────────────

@app.route("/")
@login_required
def index():
    with open(os.path.join(app.static_folder, "index.html"), encoding="utf-8") as f:
        return f.read()


@app.route("/generate", methods=["POST"])
@login_required
def generate():
    body = request.get_json(force=True, silent=True) or {}
    app.logger.info("GENERATE body keys=%s sources_len=%s", list(body.keys()), len(body.get("sources") or []))

    # Support multi-source: sources=[{type,content}, ...] or legacy single
    sources = body.get("sources", None)
    if sources is None:
        input_type = body.get("type", "url")
        content    = body.get("content", "").strip()
        if not content:
            return {"error": "No content provided."}, 400
        sources = [{"type": input_type, "content": content}]

    def stream():
        try:
            combined_parts = []
            main_title = "Report"

            for idx, src in enumerate(sources):
                src_type    = src.get("type", "url")
                src_content = src.get("content", "").strip()
                if not src_content:
                    continue

                if src_type == "url":
                    url_kind = detect_url_type(src_content)
                    if url_kind == "sheet":
                        label = f"Source {idx+1} (Google Sheet)"
                        yield sse("status", f"Fetching {label}…")
                        try:
                            title, text = fetch_google_sheet(src_content)
                        except RuntimeError as e:
                            yield sse("error", str(e))
                            return
                    else:
                        label = f"Source {idx+1} (Google Doc)"
                        yield sse("status", f"Fetching {label}…")
                        try:
                            title, text = fetch_google_doc(src_content)
                        except RuntimeError as e:
                            yield sse("error", str(e))
                            return
                    if not text:
                        yield sse("error", f"{label} appears to be empty.")
                        return
                    if idx == 0:
                        main_title = title
                    combined_parts.append(f"=== {label}: {title} ===\n{text}")
                    yield sse("status", f'Fetched: "{title}" ({len(text):,} chars)')
                else:
                    combined_parts.append(f"=== Pasted Text ===\n{src_content}")
                    if idx == 0:
                        main_title = "Pasted Report"
                    yield sse("status", f"Using pasted text ({len(src_content):,} chars)")

            if not combined_parts:
                yield sse("error", "No content provided.")
                return

            report_text = "\n\n".join(combined_parts)
            yield sse("status", "Writing case study…")

            # Keep context tight — template targets 280-350 words output
            MAX_CHARS = 3000
            if len(report_text) > MAX_CHARS:
                report_text = report_text[:MAX_CHARS] + "\n[truncated]"

            api_key = os.environ.get("ANTHROPIC_API_KEY", "").strip()
            if not api_key:
                yield sse("error", "ANTHROPIC_API_KEY is not set.")
                return

            user_message = f"Title: {main_title}\n\n{report_text}"
            payload = {
                "model": "claude-haiku-4-5-20251001",
                "max_tokens": 900,
                "system": SYSTEM_PROMPT,
                "messages": [{"role": "user", "content": user_message}],
                "stream": True,
            }
            headers = {
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
                "content-type": "application/json",
            }

            resp = requests.post(
                "https://api.anthropic.com/v1/messages",
                json=payload,
                headers=headers,
                stream=True,
                timeout=60,
            )

            if resp.status_code == 401:
                yield sse("error", "Invalid API key.")
                return
            if resp.status_code != 200:
                yield sse("error", f"API error {resp.status_code}: {resp.text[:200]}")
                return

            yield sse("writing_start", "")
            for line in resp.iter_lines():
                if not line:
                    continue
                line = line.decode("utf-8") if isinstance(line, bytes) else line
                if line.startswith("data: "):
                    data = line[6:]
                    if data == "[DONE]":
                        break
                    try:
                        chunk = json.loads(data)
                        if chunk.get("type") == "content_block_delta":
                            delta = chunk.get("delta", {})
                            if delta.get("type") == "text_delta":
                                yield sse("token", delta.get("text", ""))
                    except json.JSONDecodeError:
                        pass

            yield sse("done", "")

        except requests.exceptions.ConnectionError as e:
            yield sse("error", f"Network error reaching Anthropic API: {e}")
        except requests.exceptions.Timeout:
            yield sse("error", "Request to Anthropic API timed out.")
        except Exception as e:
            yield sse("error", f"{type(e).__name__}: {e}")

    return Response(
        stream_with_context(stream()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ── DOCX helpers ───────────────────────────────────────────────────────────────

def _shade_paragraph(paragraph, hex_color: str):
    """Apply background shading to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _set_para_border(paragraph, sides=("bottom",), color="CCCCCC", size=6, space=1):
    """Add border lines to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), str(space))
        el.set(qn("w:color"), color)
        pBdr.append(el)
    pPr.append(pBdr)


def _add_section_header(doc, text: str):
    """Add a styled section header like '01 | Participant Profile'."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(0)
    _shade_paragraph(p, "E8EDF2")
    run = p.add_run(f"  {text}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
    # Bottom border
    _set_para_border(p, sides=("bottom",), color="2563EB", size=6)


def _add_callout(doc, text: str):
    """Add a styled blockquote callout box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    _shade_paragraph(p, "EFF6FF")
    _set_para_border(p, sides=("left",), color="2563EB", size=16)
    run = p.add_run(text.strip('"').strip())
    run.italic = True
    run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
    run.font.size = Pt(11)


def _add_core_insight(doc, text: str):
    """Add the CORE INSIGHT block."""
    # Header row
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(0)
    _shade_paragraph(p, "1A3A5C")
    run = p.add_run("  CORE INSIGHT")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Body
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(6)
    _shade_paragraph(p2, "EFF6FF")
    _set_para_border(p2, sides=("left", "right", "bottom"), color="1A3A5C", size=6)
    run2 = p2.add_run(f"  {text}")
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x18)


@app.route("/download/docx", methods=["POST"])
@login_required
def download_docx():
    body    = request.get_json(force=True)
    content = body.get("content", "").strip()
    if not content:
        return {"error": "No content."}, 400

    doc = Document()

    # Default style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    lines = content.splitlines()
    in_core_insight = False
    core_insight_lines = []

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        i += 1

        # Flush core insight when we hit --- or end
        if in_core_insight:
            if line == "---" or line == "":
                if core_insight_lines:
                    _add_core_insight(doc, " ".join(core_insight_lines))
                    core_insight_lines = []
                    in_core_insight = False
                if line == "---":
                    continue
            else:
                core_insight_lines.append(line)
            continue

        if not line or line == "---":
            doc.add_paragraph()
            continue

        # H1 — person's name
        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
            p.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

        # H3 — tagline/subtitle
        elif line.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line[4:])
            run.italic = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x6B, 0x6B, 0x67)

        # H2 — numbered section header
        elif line.startswith("## "):
            _add_section_header(doc, line[3:])

        # Stats line: **stat1** · **stat2** · **stat3**
        elif re.match(r"^\*\*[^*]+\*\*\s*·", line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            parts = re.split(r"(\*\*[^*]+\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
                elif part.strip():
                    p.add_run(part)

        # Blockquote callout
        elif line.startswith("> "):
            _add_callout(doc, line[2:])

        # CORE INSIGHT marker
        elif line == "**CORE INSIGHT**":
            in_core_insight = True

        # Bold subsection header e.g. **3.1 Title** or **5.2 Title**
        elif re.match(r"^\*\*\d+\.\d+\s", line) and line.endswith("**"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            run = p.add_run(line.strip("*"))
            run.bold = True
            run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

        # Bold label with colon e.g. **The Impact:** text
        elif line.startswith("**") and ":**" in line:
            label, _, rest = line.partition(":**")
            p = doc.add_paragraph()
            p.add_run(label.lstrip("*") + ": ").bold = True
            p.add_run(rest.strip())

        # Plain bold line
        elif line.startswith("**") and line.endswith("**") and len(line) > 4:
            p = doc.add_paragraph()
            run = p.add_run(line.strip("*"))
            run.bold = True

        else:
            doc.add_paragraph(line)

    # Flush any remaining core insight at EOF
    if in_core_insight and core_insight_lines:
        _add_core_insight(doc, " ".join(core_insight_lines))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name="case_study.docx",
    )


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5001))
    app.run(host="0.0.0.0", port=port, debug=False)
